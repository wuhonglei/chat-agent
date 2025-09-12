"""Document processing service"""

import os
import uuid
from datetime import datetime
from typing import Optional

import pypdf
from docx import Document as DocxDocument
from loguru import logger
from markdown import markdown

from app.core.config import settings
from app.models.document import Document, DocumentSource, DocumentStatus
from app.services.external_fetcher import ExternalDocumentFetcher


class DocumentProcessor:
    """Process various document formats"""

    def __init__(self):
        self.external_fetcher = ExternalDocumentFetcher()

    async def process_file(
        self,
        file_path: str,
        doc_id: str,
        filename: str,
    ) -> Document:
        """Process uploaded file"""
        try:
            # Extract content based on file type
            file_extension = filename.split(".")[-1].lower()

            if file_extension == "pdf":
                content = self._extract_pdf(file_path)
            elif file_extension == "docx":
                content = self._extract_docx(file_path)
            elif file_extension == "txt":
                content = self._extract_text(file_path)
            elif file_extension == "md":
                content = self._extract_markdown(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Create document
            document = Document(
                id=doc_id,
                name=filename,
                source=DocumentSource.LOCAL,
                content=content,
                metadata={
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
                status=DocumentStatus.COMPLETED,
                created_at=datetime.now(),
                updated_at=datetime.now(),
            )

            return document

        except Exception as e:
            logger.error(f"Failed to process file {filename}: {e}")
            raise

    async def process_url(
        self,
        url: str,
        source: DocumentSource,
    ) -> Document:
        """Process external document from URL"""
        try:
            doc_id = str(uuid.uuid4())

            # Fetch content based on source
            if source == DocumentSource.CONFLUENCE:
                content, metadata = await self.external_fetcher.fetch_confluence(url)
                name = metadata.get("title", "Confluence Document")
            elif source == DocumentSource.GOOGLE_DOCS:
                content, metadata = await self.external_fetcher.fetch_google_docs(url)
                name = metadata.get("title", "Google Document")
            elif source == DocumentSource.GOOGLE_SLIDES:
                content, metadata = await self.external_fetcher.fetch_google_slides(url)
                name = metadata.get("title", "Google Slides")
            else:
                raise ValueError(f"Unsupported external source: {source}")

            # Create document
            document = Document(
                id=doc_id,
                name=name,
                source=source,
                source_url=url,
                content=content,
                metadata=metadata,
                status=DocumentStatus.COMPLETED,
                created_at=datetime.now(),
                updated_at=datetime.now(),
            )

            return document

        except Exception as e:
            logger.error(f"Failed to process URL {url}: {e}")
            raise

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = pypdf.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text

    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from markdown file"""
        with open(file_path, "r", encoding="utf-8") as file:
            md_content = file.read()
            # Convert to plain text
            from bs4 import BeautifulSoup

            html = markdown(md_content)
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text()
